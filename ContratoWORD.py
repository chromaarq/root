import tkinter as tk
from tkinter import messagebox, font, filedialog
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt, Inches 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from docx.oxml.ns import qn 
from docx.oxml import OxmlElement
from PIL import Image, ImageTk
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

beje = '#e4d8ce'
verde = '#394731'
cinza = '#d0cfcc'
marrom = '#8f3c23'



# Função para ler o template do contrato de um arquivo 
def ler_template_contrato(): 
    with open('template_contrato.txt', 'r', encoding='utf-8') as file: 
        template = file.read() 
    return template 
# Carrega o template do contrato 
template_contrato = ler_template_contrato()

# Função para gerar o contrato
def gerar_contrato(dados):
    dia = 0; mes = 0; ano = 0
    d = dados['data_contrato']
    i = 0

    for palavra in d.split('/'):
        if i==0:
            dia = palavra
        if i==1:
            if palavra=="01":
                mes = "Janeiro"
            if palavra=="02":
                mes = "Fevereiro"
            if palavra=="03":
                mes = "Março"
            if palavra=="04":
                mes = "Abril"
            if palavra=="05":
                mes = "Maio"
            if palavra=="06":
                mes = "Junho"
            if palavra=="07":
                mes = "Julho"
            if palavra=="08":
                mes = "Agosto"
            if palavra=="09":
                mes = "Setembro"
            if palavra=="10":
                mes = "Outubro"
            if palavra=="11":
                mes = "Novembro"
            if palavra=="12":
                mes = "Dezembro"
        if i==2:
            ano = palavra
        i=i+1
    
    print(dia)
    print(mes)
    print(ano)

    contrato = template_contrato.format(
        nome_cliente=dados['nome_cliente'],
        endereco_cliente=dados['endereco_cliente'],
        rg_cliente=dados['rg_cliente'],
        cpf_cliente=dados['cpf_cliente'],
        telefone_cliente=dados['telefone_cliente'],
        email_cliente=dados['email_cliente'],
        descricao_servico=dados['descricao_servico'],
        area_projeto=dados['area_projeto'],
        endereco_imovel=dados['endereco_imovel'],
        unidade_imovel=dados['unidade_imovel'],
        condominio_imovel=dados['condominio_imovel'],
        bairro_imovel=dados['bairro_imovel'],
        cidade_imovel=dados['cidade_imovel'],
        valor_servico=dados['valor_servico'],
        forma_pagamento=dados['forma_pagamento'],
        dia_contrato=dia,
        mes_contrato=mes,
        ano_contrato=ano,
    )
    return contrato

# Função para coletar os dados da interface gráfica
def coletar_dados():
    dados = {
        'nome_cliente': entry_nome_cliente.get(),
        'endereco_cliente': entry_endereco_cliente.get(),
        'rg_cliente': entry_rg_cliente.get(),
        'cpf_cliente': entry_cpf_cliente.get(),
        'telefone_cliente': entry_telefone_cliente.get(),
        'email_cliente': entry_email_cliente.get(),
        'descricao_servico': entry_descricao_servico.get("1.0", tk.END).strip(),
        'area_projeto': entry_area_projeto.get(),
        'endereco_imovel': entry_endereco_imovel.get(),
        'unidade_imovel': entry_unidade_imovel.get(),
        'condominio_imovel': entry_condominio_imovel.get(),
        'bairro_imovel': entry_bairro_imovel.get(),
        'cidade_imovel': entry_cidade_imovel.get(),
        'valor_servico': entry_valor_servico.get(),
        'forma_pagamento': entry_forma_pagamento.get("1.0", tk.END).strip(),
        'data_contrato': entry_data_contrato.get()
    }
        
    contrato = gerar_contrato(dados)
    exibir_contrato(contrato)

# Função para exibir o contrato em uma nova janela e salvar em PDF ou DOCX
def exibir_contrato(contrato):
    janela_contrato = tk.Toplevel(root)
    janela_contrato.title("Contrato Gerado")

    texto_contrato = tk.Text(janela_contrato, wrap='word', font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))
    texto_contrato.insert('1.0', contrato)
    texto_contrato.pack(expand=True, fill='both')
    
    def salvar_pdf():
        arquivo_pdf = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if arquivo_pdf:
            c = canvas.Canvas(arquivo_pdf, pagesize=letter)
            largura, altura = letter

            # Adicionar logomarca
            if logomarca_path:
                c.drawImage(logomarca_path, 50, altura - 100, width=100, height=50)

            # Adicionar texto do contrato
            c.drawString(50, altura - 150, contrato)
            c.save()
            messagebox.showinfo("Salvo em PDF", f"Contrato salvo como {arquivo_pdf}")
    
    def salvar_docx(): 
        arquivo_docx = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")]) 
        if arquivo_docx: 
            doc = Document() 
        # Configurar as margens do documento 
        section = doc.sections[0] 
        section.left_margin = Inches(1) 
        section.right_margin = Inches(1) 
        section.top_margin = Inches(1) 
        section.bottom_margin = Inches(1)

        # Adicionar a logomarca no cabeçalho 
        if logomarca_path: 
            section = doc.sections[0] 
            header = section.header 
            header_paragraph = header.paragraphs[0] 
            run = header_paragraph.add_run() 
            run.add_picture(logomarca_path, width=Pt(120), height=Pt(50)) 
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 

        # # Adicionar o texto do contrato 
        # p = doc.add_paragraph(contrato) 
        # Adicionar o texto do contrato com negrito em partes específicas 
        for linha in contrato.split('\n'): 
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            if "CLÁUSULA" in linha or "CONTRATO:" in linha: 
                # p = doc.add_paragraph() 
                run = p.add_run(linha) 
                run.bold = True 
            else:
                for palavra in linha.split('-'):
                    if ("1." in palavra or "2." in palavra or "3." in palavra or "4." in palavra or "5." in palavra or "6." in palavra or "7." in palavra or "8." in palavra or "9." in palavra or "10." in palavra or "11." in palavra or "12." in palavra) and ("CNPJ" not in palavra) and ("RG" not in palavra) and ("etapa" not in palavra) and ("parágrafo" not in palavra) and ("cláusula" not in palavra) and ("Lei" not in palavra):
                        # p = doc.add_paragraph()
                        run = p.add_run(palavra)
                        run.bold = True 
                        run = p.add_run("-")
                        run.bold = True 
                    else:
                        # p = doc.add_paragraph() 
                        run = p.add_run(palavra)
                    if 'bold' in estilo_fonte_var.get(): 
                        run.bold = True 
                    if 'italic' in estilo_fonte_var.get(): 
                        run.italic = True

        # Configurar espaçamento entre linhas
        

        # Configurar a fonte e tamanho 
        style = doc.styles['Normal'] 
        font = style.font 
        font.name = fonte_var.get() 
        font.size = Pt(tamanho_fonte_var.get()) 
        
        for run in p.runs: 
            run.font.name = fonte_var.get() 
            run.font.size = Pt(tamanho_fonte_var.get()) 
            if 'bold' in estilo_fonte_var.get(): 
                run.bold = True 
            if 'italic' in estilo_fonte_var.get(): 
                run.italic = True 
        
        doc.save(arquivo_docx) 
        messagebox.showinfo("Salvo em DOCX", f"Contrato salvo como {arquivo_docx}")
    
    botao_salvar_pdf = tk.Button(janela_contrato, text="Salvar como PDF", command=salvar_pdf)
    botao_salvar_pdf.pack(pady=10)

    botao_salvar_docx = tk.Button(janela_contrato, text="Salvar como DOCX", command=salvar_docx)
    botao_salvar_docx.pack(pady=10)

# Função para carregar a logomarca
# def carregar_logomarca():
#     global logomarca_path
#     logomarca_path = filedialog.askopenfilename(title="Selecione a Logomarca", filetypes=[("Image Files", "*.png;*.jpg;*.jpeg")])
#     if logomarca_path:
#         label_logomarca.config(text=f"Logomarca Carregada: {logomarca_path}")


# Variável global para o caminho da logomarca
logomarca_path = "Logotipo_Chroma.png"

# # Mostra a Logomarca
# root = tk.Tk()
# root.title("Exibição da logomarca")
# # Obter tamanho da imagem
# with Image.open('Logotipo_Chroma.png') as img:
#     width, height = img.size
#     width = int(width*0.1)
#     height = int(height*0.1)
#     resized_image = img.resize((width,height))
# # root.geometry("800x600")
# photo = ImageTk.PhotoImage(resized_image)
# # Criar um widget Label para exibir a imagem
# label = tk.Label(root, image=photo)
# label.image = photo  # Necessário para manter a referência
# label.pack()


# Configuração da Interface Gráfica
root = tk.Tk()
fonte_var = tk.StringVar(value="Nunito")
estilo_fonte_var = tk.StringVar(value="normal")
tamanho_fonte_var = tk.IntVar(value=11)
root.title("Gerador de Contratos")
root.config(bg=cinza)

tk.Label(root, text="Nome do Cliente:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=0, column=0)
entry_nome_cliente = tk.Entry(root,bg=cinza)
entry_nome_cliente.grid(row=0, column=1)
entry_nome_cliente.config(width=50,justify="left",font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Label(root, text="Endereço do Cliente:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=1, column=0)
entry_endereco_cliente = tk.Entry(root,bg=cinza)
entry_endereco_cliente.grid(row=1, column=1)
entry_endereco_cliente.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Label(root, text="RG do Cliente:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=2, column=0)
entry_rg_cliente = tk.Entry(root,bg=cinza)
entry_rg_cliente.grid(row=2, column=1)
entry_rg_cliente.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Label(root, text="CPF Cliente:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=3, column=0)
entry_cpf_cliente = tk.Entry(root,bg=cinza)
entry_cpf_cliente.grid(row=3, column=1)
entry_cpf_cliente.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Label(root, text="Telefone do Cliente:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=4, column=0)
entry_telefone_cliente = tk.Entry(root,bg=cinza)
entry_telefone_cliente.grid(row=4, column=1)
entry_telefone_cliente.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.LabelFrame(root, text="e-mail do Cliente:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=5, column=0)
entry_email_cliente = tk.Entry(root,bg=cinza)
entry_email_cliente.grid(row=5, column=1)
entry_email_cliente.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Label(root, text="Área do Protjeto:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=0, column=2)
entry_area_projeto = tk.Entry(root,bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))
entry_area_projeto.grid(row=0, column=3)
entry_area_projeto.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Label(root, text="Endereço do Imóvel",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=1, column=2)
entry_endereco_imovel = tk.Entry(root,bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))
entry_endereco_imovel.grid(row=1, column=3)
entry_endereco_imovel.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Label(root, text="Unidade do Imóvel:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=2, column=2)
entry_unidade_imovel = tk.Entry(root,bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))
entry_unidade_imovel.grid(row=2, column=3)
entry_unidade_imovel.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Label(root, text="Condomínio do Imóvel:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=3, column=2)
entry_condominio_imovel = tk.Entry(root,bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))
entry_condominio_imovel.grid(row=3, column=3)
entry_condominio_imovel.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Label(root, text="Bairro do Imóvel:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=4, column=2)
entry_bairro_imovel = tk.Entry(root,bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))
entry_bairro_imovel.grid(row=4, column=3)
entry_bairro_imovel.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Label(root, text="Cidade do Imóvel:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=5, column=2)
entry_cidade_imovel = tk.Entry(root,bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))
entry_cidade_imovel.grid(row=5, column=3)
entry_cidade_imovel.config(width=50,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))


tk.Label(root, text="Descrição do Serviço:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=6, column=0)
entry_descricao_servico = tk.Text(root,bg=cinza)
entry_descricao_servico.grid(row=6, column=1, rowspan=6)
entry_descricao_servico.config(width=50, height=20,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))



tk.Label(root, text="Valor do Serviço:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=6, column=2)
entry_valor_servico = tk.Entry(root,bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))
entry_valor_servico.grid(row=6, column=3)
entry_valor_servico.config(width=50)


tk.Label(root, text="Forma de Pagamento:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=7, column=2)
entry_forma_pagamento = tk.Text(root,bg=cinza)
entry_forma_pagamento.grid(row=7, column=3)
entry_forma_pagamento.config(width=50, height=10,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))


tk.Label(root, text="Data do Contrato:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=8, column=2)
entry_data_contrato = tk.Entry(root,bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))
entry_data_contrato.grid(row=8, column=3)
entry_data_contrato.config(width=50)

# Seleção de Fonte
tk.Label(root, text="Fonte:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=9, column=2)
tk.OptionMenu(root, fonte_var,"Nunito", "Arial", "Helvetica", "Times", "Courier", "Verdana").grid(row=9, column=3)
option_menu_fonte = tk.OptionMenu(root, fonte_var,"Nunito", "Arial", "Helvetica", "Times", "Courier", "Verdana")
option_menu_fonte.config(bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

# Seleção de Tamanho de Fonte
tk.Label(root, text="Tamanho da Fonte:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=10, column=2)
# tamanho_fonte_var = tk.IntVar(value=10)
tk.Spinbox(root, from_=8, to=72, textvariable=tamanho_fonte_var,bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()),width=5).grid(row=10, column=3)

## Seleção de Estilo de Fonte
tk.Label(root, text="Estilo da Fonte:",bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get())).grid(row=11, column=2)
# estilo_fonte_var = tk.StringVar(value="normal")
tk.OptionMenu(root, estilo_fonte_var, "normal", "bold", "italic", "bold italic").grid(row=11, column=3)
option_menu_estilo = tk.OptionMenu(root, estilo_fonte_var, "normal", "bold", "italic", "bold italic")
option_menu_estilo.config(bg=cinza,font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))

tk.Button(root, text="Gerar Contrato", command=coletar_dados,width=15, height=5, bg=beje).grid(row=15, column=3)
botao_contrato = tk.Button(root)
botao_contrato.config(font=(fonte_var.get(), tamanho_fonte_var.get(), estilo_fonte_var.get()))
# tk.Button(root, text="Carregar Logo", command=carregar_logomarca).grid(row=16, columnspan=2)
# label_logomarca = tk.Label(root, text="Logomarca não carregada") 
# label_logomarca.config(text=f"Logomarca Carregada: {logomarca_path}")

# Abrir imagem da logo e redimensionar
with Image.open('Logotipo_Chroma.png') as img:
    width, height = img.size
    width = int(width*0.2)
    height = int(height*0.2)
    resized_image = img.resize((width,height))

# Criar um widget Label para exibir a imagem da logo
image = Image.open('Logotipo_Chroma.png')
photo = ImageTk.PhotoImage(resized_image)
label = tk.Label(root, image=photo)
label.config(justify="left",bg=cinza)
tk.Label(root, image=photo,bg=cinza).grid(row=15, column=0)


root.mainloop()